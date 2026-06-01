from app.engines.knowledge_engine.parser.base import BaseDocumentParser


class WordParser(BaseDocumentParser):
    SUPPORTED_EXTENSIONS = [".docx"]

    def parse(self, file_path: str, **kwargs) -> dict:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            tables.append(table_data)
        return {
            "content": "\n".join(paragraphs),
            "tables": tables,
            "metadata": {"format": "docx", "paragraphs": len(paragraphs), "tables": len(tables)},
        }
